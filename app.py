import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from groq import Groq
import os
from docx import Document
from io import BytesIO

# --- Streamlit Page Setup ---
st.set_page_config(page_title="Smart ICU Health Monitor", layout="wide")
st.title("🏥 Smart ICU Health Monitoring Dashboard")

# --- Sidebar for API Key ---
st.sidebar.header("🔑 API Key Setup")
if "api_key" not in st.session_state:
    st.session_state.api_key = ""

api_key_input = st.sidebar.text_input("Enter your GROQ API Key:", type="password")
submit_key = st.sidebar.button("✅ Submit API Key")

if submit_key:
    if not api_key_input.strip():
        st.sidebar.error("⚠️ Please enter a valid API key.")
    else:
        st.session_state.api_key = api_key_input.strip()
        st.sidebar.success("✅ API key saved successfully!")

# Sidebar Overall Status checkbox (below API key)

# Stop if API key missing
if not st.session_state.api_key:
    st.warning("⚠️ Please enter your GROQ API key in the sidebar to continue.")
    st.stop()

# Sidebar Overall Status controls (always visible)
if "show_overall" not in st.session_state:
    st.session_state.show_overall = False
if "overall_scope" not in st.session_state:
    st.session_state.overall_scope = "Current Sheet"

overall_scope = st.sidebar.selectbox("Overall Scope", options=["Current Sheet", "All Sheets"], index=0)
if st.sidebar.button("Show Overall Status"):
    st.session_state.show_overall = True
    st.session_state.overall_scope = overall_scope
if st.sidebar.button("Hide Overall Status"):
    st.session_state.show_overall = False

# --- Initialize Groq Client ---
try:
    client = Groq(api_key=st.session_state.api_key)
except Exception as e:
    st.error(f"❌ Failed to initialize Groq client: {e}")
    st.stop()

# --- File Upload ---
uploaded_file = st.file_uploader("📤 Upload Patient Data (Excel file)", type=["xlsx"])

# --- Flexible Column Detection ---
expected_columns = {
    "patient_name": ["patient", "name", "patient_name"],
    "blood_pressure(mmhg)": ["bp", "blood pressure", "pressure", "blood_pressure"],
    "temperature(°c)": ["temp", "temperature", "body temp"],
    "oxygen_level(%)": ["spo2", "oxygen", "oxygen level", "oxygen saturation"],
    "urine_level(ml)": ["urine", "urine output", "urine_level"],
    "urine_ph": ["urine ph", "ph", "urine acidity"],
    "blood_count(million/ul)": ["blood count", "rbc", "blood cells", "blood_count"]
}

def auto_map_columns(df):
    mapping = {}
    for expected, variants in expected_columns.items():
        for col in df.columns:
            for variant in variants:
                if variant.lower() in col.lower():
                    mapping[col] = expected
    return df.rename(columns=mapping)


def classify_row(r):
    try:
        oxy = pd.to_numeric(r.get('oxygen_level(%)'), errors='coerce')
    except Exception:
        oxy = None
    try:
        bp = pd.to_numeric(r.get('blood_pressure(mmhg)'), errors='coerce')
    except Exception:
        bp = None
    try:
        temp = pd.to_numeric(r.get('temperature(°c)'), errors='coerce')
    except Exception:
        temp = None

    # Prioritize dangerous > serious > normal
    if oxy is not None and oxy < 85:
        return 'Dangerous'
    if bp is not None and bp > 180:
        return 'Dangerous'
    if temp is not None and temp >= 40:
        return 'Dangerous'

    if oxy is not None and oxy < 90:
        return 'Serious'
    if bp is not None and bp > 140:
        return 'Serious'
    if temp is not None and temp >= 39:
        return 'Serious'

    return 'Normal'

# --- Process Uploaded File ---
if uploaded_file is not None:
    try:
        # Read all sheets so we can show either current or all-sheet overview
        sheets = pd.read_excel(uploaded_file, sheet_name=None)
        sheet_names = list(sheets.keys())
        sheet_choice = st.selectbox("📑 Select sheet to load", options=["All Sheets"] + sheet_names)

        # build df_all (concatenated) and df (selected)
        df_list = []
        for sname, sdf in sheets.items():
            sdf2 = sdf.copy()
            sdf2["_sheet_name"] = sname
            df_list.append(sdf2)
        df_all = pd.concat(df_list, ignore_index=True, sort=False) if df_list else pd.DataFrame()

        if sheet_choice == "All Sheets":
            df = df_all.copy()
        else:
            df = sheets.get(sheet_choice, pd.DataFrame()).copy()

        # map columns on both current df and full df
        df = auto_map_columns(df)
        df_all = auto_map_columns(df_all)

        # Ensure all columns exist
        for col in expected_columns.keys():
            if col not in df.columns:
                df[col] = None

        st.success("✅ File uploaded and columns mapped successfully!")

        # If the user requested overall status, show a small pie in the sidebar immediately
        if st.session_state.get('show_overall', False):
            try:
                scope = st.session_state.get('overall_scope', 'Current Sheet')
                data_for_ward = df_all if scope == 'All Sheets' else df
                ward_categories_sb = data_for_ward.apply(classify_row, axis=1)
                counts_sb = ward_categories_sb.value_counts()
                # Merge 'Dangerous' into 'Serious' for a two-category view
                serious_sb = int(counts_sb.get('Dangerous', 0)) + int(counts_sb.get('Serious', 0))
                normal_sb = int(counts_sb.get('Normal', 0))
                labels_sb = ['Serious', 'Normal']
                counts_list_sb = [serious_sb, normal_sb]
                fig_sb, ax_sb = plt.subplots(figsize=(3, 3))
                if sum(counts_list_sb) == 0:
                    ax_sb.text(0.5, 0.5, 'No data', ha='center', va='center')
                else:
                    colors = ['#ff7f0e', '#2ca02c']
                    ax_sb.pie(counts_list_sb, labels=labels_sb, colors=colors, startangle=90,
                              autopct=lambda pct: f"{pct:.1f}% ({int(pct*sum(counts_list_sb)/100)})")
                    ax_sb.set_title('Ward Risk (Sidebar)')
                fig_sb.tight_layout()
                st.sidebar.pyplot(fig_sb)
                plt.close(fig_sb)
            except Exception:
                pass

    except Exception as e:
        st.error(f"❌ Error reading file: {e}")
        st.stop()

    # --- Overview ---
        # --- Full Dataset Overview Section ---
        st.subheader("🌐 Full Dataset Overview")
        st.write("Loaded sheet:", sheet_choice)
        st.dataframe(df)

        # --- Dataset-level AI Analysis ---
        st.subheader("🤖 Dataset-level AI Summary")
        if st.button("🧾 Generate Dataset Report"):
            with st.spinner("Analyzing dataset..."):
                try:
                    # Build simple numeric summary for prompt
                    numeric_cols = [
                        "blood_pressure(mmhg)", "temperature(°c)", "oxygen_level(%)",
                        "urine_level(ml)", "urine_ph", "blood_count(million/ul)"
                    ]
                    stats = {}
                    for c in numeric_cols:
                        if c in df.columns:
                            try:
                                col_series = pd.to_numeric(df[c], errors='coerce')
                                stats[c] = {
                                    "count": int(col_series.count()),
                                    "mean": float(col_series.mean(skipna=True)) if col_series.count() > 0 else None,
                                    "min": float(col_series.min(skipna=True)) if col_series.count() > 0 else None,
                                    "max": float(col_series.max(skipna=True)) if col_series.count() > 0 else None
                                }
                            except Exception:
                                stats[c] = None

                    ai_prompt = f"""
                    You are an expert ICU medical AI.
                    Provide an executive summary and insights based on the following dataset-level statistics.

                    Loaded Sheet: {sheet_choice}
                    Number of rows: {len(df)}
                    Number of unique patients: {df.get('patient_name', pd.Series()).nunique()}
                    Column stats: {stats}

                    Please:
                    - Summarize overall patient status trends.
                    - Highlight any concerning aggregate values or anomalies.
                    - Recommend next steps for clinical review and monitoring.
                    """

                    completion = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[
                            {"role": "system", "content": "You are a skilled ICU data analyst."},
                            {"role": "user", "content": ai_prompt}
                        ],
                        temperature=0.5,
                        max_tokens=400
                    )
                    result = completion.choices[0].message.content
                    st.success("✅ Dataset AI Report")
                    st.write(result)

                    # Word report
                    doc = Document()
                    doc.add_heading(f"Dataset Report: {sheet_choice}", level=1)
                    doc.add_paragraph(f"Rows: {len(df)}")
                    doc.add_paragraph(f"Unique patients: {df.get('patient_name', pd.Series()).nunique()}")
                    doc.add_heading("Column Statistics", level=2)
                    for k, v in stats.items():
                        doc.add_paragraph(f"{k}: {v}")
                    doc.add_heading("AI Analysis", level=2)
                    doc.add_paragraph(result)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="📥 Download Dataset Report",
                        data=buffer,
                        file_name=f"Dataset_Report_{sheet_choice}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"❌ Error while generating dataset AI report: {e}")

        # --- Ward-level Health Overview (shown when sidebar 'Overall Status' is checked) ---
        if st.session_state.get('overall_status', False):
            st.subheader("🏥 Ward-level Health Overview")
            try:
                # choose data based on sidebar scope
                scope = st.session_state.get('overall_scope', 'Current Sheet')
                if scope == 'All Sheets':
                    data_for_ward = df_all if 'df_all' in locals() else df
                else:
                    data_for_ward = df

                ward_categories = data_for_ward.apply(classify_row, axis=1)
                counts = ward_categories.value_counts()
                # Ensure consistent order
                labels = ['Dangerous', 'Serious', 'Normal']
                counts_list = [int(counts.get(l, 0)) for l in labels]
                total = sum(counts_list)

                fig, ax = plt.subplots(figsize=(4, 4))
                # show raw counts for debugging/visibility
                st.write("Ward counts:", {labels[i]: counts_list[i] for i in range(len(labels))})
                if total == 0:
                    ax.text(0.5, 0.5, 'No data', ha='center', va='center')
                else:
                    colors = ['#d62728', '#ff7f0e', '#2ca02c']
                    wedges, texts, autotexts = ax.pie(counts_list, labels=labels, autopct=lambda pct: f"{pct:.1f}% ({int(pct*total/100)})",
                                                      colors=colors, startangle=90)
                    ax.set_title('Ward Risk Distribution')
                    ax.legend(wedges, labels, title="Category", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
                    fig.tight_layout()

                fig.tight_layout()
                fig.tight_layout()
                st.pyplot(fig)
                plt.close(fig)
                plt.close(fig)

                # Ask AI for a short ward-level interpretation
                ai_prompt_ward = f"""
                You are an expert ICU medical AI.
                Based on the ward counts: Dangerous={counts_list[0]}, Serious={counts_list[1]}, Normal={counts_list[2]},
                and total patients={total}, provide:
                - A concise interpretation of the ward's current risk level.
                - Recommended actions for ward triage and monitoring.
                - If applicable, thresholds to escalate to higher care.
                """

                if total > 0:
                    try:
                        completion = client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=[
                                {"role": "system", "content": "You are a skilled ICU data analyst."},
                                {"role": "user", "content": ai_prompt_ward}
                            ],
                            temperature=0.4,
                            max_tokens=250
                        )
                        ward_result = completion.choices[0].message.content
                        st.markdown("**AI Ward Interpretation**")
                        st.write(ward_result)
                    except Exception as e:
                        st.error(f"❌ Error generating ward AI summary: {e}")
                else:
                    st.info("No patient rows to analyze for ward status.")

            except Exception as e:
                st.error(f"❌ Error computing ward overview: {e}")

    # --- Patient Selection ---
    st.subheader("🩺 Individual Patient Health Analysis")
    if "patient_name" in df.columns:
        patient_col = df["patient_name"]
        # If multiple columns were mapped to the same name, pandas returns a DataFrame
        if isinstance(patient_col, pd.DataFrame):
            patient_series = patient_col.apply(lambda row: next((x for x in row if pd.notna(x)), None), axis=1)
        else:
            patient_series = patient_col

        unique_patients = patient_series.dropna().unique().tolist()
        if len(unique_patients) == 0:
            st.info("No patient names found in `patient_name` column.")
            patient_data = pd.Series({k: None for k in expected_columns.keys()})
            selected_patient = None
        else:
            selected_patient = st.selectbox("Select a Patient", unique_patients)
            match_idx = patient_series[patient_series == selected_patient].index
            if len(match_idx) == 0:
                st.error("Selected patient not found in rows.")
                patient_data = pd.Series({k: None for k in expected_columns.keys()})
            else:
                patient_data = df.loc[match_idx[0]]
    else:
        st.info("No `patient_name` column found in the dataset.")
        patient_data = pd.Series({k: None for k in expected_columns.keys()})

    metrics = {
        "Blood Pressure (mmHg)": patient_data["blood_pressure(mmhg)"],
        "Temperature (°C)": patient_data["temperature(°c)"],
        "Oxygen Level (%)": patient_data["oxygen_level(%)"],
        "Urine Level (ml)": patient_data["urine_level(ml)"],
        "Urine pH": patient_data["urine_ph"],
        "Blood Count (million/μL)": patient_data["blood_count(million/ul)"]
    }

    # --- Graph Section ---
    st.markdown("### 📈 Patient Health Metrics Visualization")

    # Show small graphs side by side (2 rows, 3 per row)
    metric_items = list(metrics.items())
    for i in range(0, len(metric_items), 3):
        row = metric_items[i:i+3]
        cols = st.columns(len(row))
        for j, (metric_name, value) in enumerate(row):
            with cols[j]:
                fig, ax = plt.subplots(figsize=(2.8, 2.8))
                # Safely handle missing or non-numeric values
                try:
                    if pd.isna(value):
                        raise ValueError("No data")
                    val = float(value)
                    ax.bar([metric_name], [val], color="royalblue")
                    # set a reasonable lower bound for visibility
                    if val >= 0:
                        ax.set_ylim(bottom=0)
                except Exception:
                    ax.text(0.5, 0.5, "No data", ha='center', va='center', fontsize=10, color='gray')
                    ax.set_ylim(0, 1)

                ax.set_title(metric_name, fontsize=9)
                ax.set_ylabel("Value", fontsize=8)
                ax.tick_params(axis='x', rotation=45, labelsize=7)
                ax.grid(axis="y", linestyle="--", alpha=0.5)
                st.pyplot(fig)

    # --- AI Analysis for Selected Patient ---
    st.subheader("🤖 AI Health Analysis Report")
    if st.button("🧠 Generate Patient Report"):
        with st.spinner(f"Analyzing health metrics for {selected_patient}..."):
            try:
                ai_prompt = f"""
                You are an expert ICU medical AI.
                Analyze this patient's vital signs and provide:
                - Interpretation of each parameter
                - Any potential health risks
                - Recommended next steps for doctors or nurses

                Patient Name: {selected_patient}
                Health Data: {metrics}
                """

                completion = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",  # ✅ Updated & supported model
                    messages=[
                        {"role": "system", "content": "You are a skilled ICU medical analyst."},
                        {"role": "user", "content": ai_prompt}
                    ],
                    temperature=0.5,
                    max_tokens=300
                )
                result = completion.choices[0].message.content
                st.success(f"✅ AI Report for {selected_patient}")
                st.write(result)

                # --- Generate Word Report ---
                doc = Document()
                doc.add_heading(f"Patient Health Report: {selected_patient}", level=1)
                doc.add_paragraph("Here are the recorded health metrics and AI analysis for the patient.\n")

                for metric, value in metrics.items():
                    doc.add_paragraph(f"{metric}: {value}")

                doc.add_heading("AI Analysis", level=2)
                doc.add_paragraph(result)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📥 Download Word Report",
                    data=buffer,
                    file_name=f"{selected_patient}_Health_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"❌ Error while generating AI report: {e}")

else:
    st.info("📂 Please upload a valid Excel file to continue.")
